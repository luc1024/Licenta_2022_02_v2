import os

# "pip install python-docx==0.8.10": https://automatetheboringstuff.com/2e/chapter15/ ; functioneaza si 0.8.11
import docx

from docx.document import Document  # https://stackoverflow.com/a/52466366/2358837


def write_assignment_item(index_item, item, doc, verify=False):
    if len(doc.paragraphs) == 1 and len(doc.paragraphs[0].runs) == 0:
        para = doc.paragraphs[0]
    else:
        doc.add_paragraph()
        doc.add_paragraph()
        para = doc.paragraphs[-1]

    category_course = '[C' + str(item['category']) + '/M' + str(item['course']) \
                + '/I' + str(item['original_index']+1) + '] ' \
                if verify else ''

    para.text = category_course + str(index_item + 1) + '. ' + item['question']  # intrebarea
    para.runs[0].italic = True

    num = ['a) ', 'b) ', 'c) ', 'd) ']
    for i in range(4):
        correct = '[' + str(item['answers'][i]['correct']) + ']' if verify else ''
        doc.add_paragraph(correct + num[i] + item['answers'][i]['text'])
    return doc


def write_answer_item(index_item, item, doc):
    if len(doc.paragraphs) == 1 and len(doc.paragraphs[0].runs) == 0:
        para = doc.paragraphs[0]
    else:
        para = doc.paragraphs[-1]
    para.add_run(str(index_item + 1) + '.')  # intrebarea
    num = ['a) ', 'b) ', 'c) ', 'd) ']
    for i in range(4):
        if item['answers'][i]['correct']:
            para.add_run(num[i])
            para.runs[-1].bold = True
    return doc


def write_assignment(path: str, filename: str, assignment: dict, verify=False):
    # https://stackoverflow.com/a/61822452/2358837
    doc: docx.document.Document = docx.Document('templates/Template.docx')
    for index_item, item in enumerate(assignment['items']):
        write_assignment_item(index_item, item, doc, verify)
    doc.save(path + filename)
    print('Doc created ' + filename)


def write_solution(path: str, filename: str, assignment: dict):
    # https://stackoverflow.com/a/61822452/2358837
    doc: docx.document.Document = docx.Document('templates/Template.docx')
    for index_item, item in enumerate(assignment['items']):
        write_answer_item(index_item, item, doc)
    doc.save(path + filename)
    print('Doc created ' + filename)


def write_docs(assignments: list, doc_folder, verify=False):
    os.mkdir(doc_folder)
    os.mkdir(doc_folder + 'assignments')
    os.mkdir(doc_folder + 'solutions')
    for assignment in assignments:
        write_assignment(doc_folder + 'assignments' + os.sep,
                      'Assignment_%s%s_%02d.docx' % (assignment['day'], assignment['exam_board'], assignment['index'] + 1),
                         assignment, verify=verify)
        write_solution(doc_folder + 'solutions' + os.sep,
                      'Solution_%s%s_%02d.docx' % (assignment['day'], assignment['exam_board'], assignment['index'] + 1),
                       assignment)
