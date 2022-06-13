import os

import docx  # "pip install python-docx==0.8.10": https://automatetheboringstuff.com/2e/chapter15/
from docx.document import Document  # https://stackoverflow.com/a/52466366/2358837

from prepare_items import get_ord_spreadsheet_values
from env_class import Env

env = Env()


def write_items_list(sheet_id, json_keyfile_name, output_dir_name):
    ord_courses_values = get_ord_spreadsheet_values(sheet_id, json_keyfile_name)

    nr_items = []
    for row in ord_courses_values:
        row = list(filter(None, row))  # elimina ''
        nr_items.append((len(row) - env.offset_col) // env.nr_paragraphs_per_item)
    num = ['a) ', 'b) ', 'c) ', 'd) ']
    for i in range(env.courses_count):
        doc: docx.document.Document = docx.Document(
            'templates/Template.docx')  # https://stackoverflow.com/a/61822452/2358837
        for j in range(nr_items[i]):
            for k in range(env.nr_paragraphs_per_item):
                pre = str(j + 1) + '. ' if k == 0 else num[k - 1]
                text = pre + ord_courses_values[i][env.offset_col + j * env.nr_paragraphs_per_item + k].strip()
                if j == 0 and k == 0:
                    doc.paragraphs[0].text = text
                else:
                    doc.add_paragraph(text)
                doc.paragraphs[-1].runs[0].italic = (k == 0)  # intrebarea cu italic
                doc.paragraphs[-1].runs[0].bold = (k == 1)  # raspunsul corect cu bold
            doc.add_paragraph('')

            course_no, materia = ord_courses_values[i][3].split('.')
            email = ord_courses_values[i][1]
            file_name = ('%02d.' % int(course_no)) + materia + ' [' + email + ']' + ' (' + str(nr_items[i]) + ' itemi)'
            doc.save(output_dir_name + os.sep + 'Liste_Itemi/' + file_name + '.docx')
    print('env.courses_count: ', env.courses_count)
